# -*- coding: utf-8 -*-
"""
Created on Sat Oct  5 02:04:28 2024

@author: phili
"""

from docx import Document

def create_email(greeting = 'Greetings,',
sender = 'EXAMPLE ESTIEMER',
company_name = 'INSERT COMPANY NAME',
extra_reasons = '(ADD ONE OR TWO MORE REASONS)'):

    document = Document()
    
    document.add_heading('Email Approach', 0)

    document.add_paragraph().add_run('Subject: ESTIEM Collaboration opportunities - Partnership offer').bold = True
    
    document.add_paragraph(greeting)
    
    p = document.add_paragraph('My name is ' + sender + ' and I am contacting you in the name of a student organisation called ')
    r = p.add_run('ESTIEM')
    r.bold = True
    r.italic = True
    p.add_run(' (estiem.org), the biggest European-wide student organisation for ')
    p.add_run('I').bold = True
    p.add_run('ndustrial ')
    p.add_run('E').bold = True
    p.add_run('ngineering and ')
    p.add_run('M').bold = True
    p.add_run('anagement students with approximately ')
    p.add_run('70 000 students').bold = True
    p.add_run(' reached (among them, ')
    p.add_run('10 000').bold = True
    p.add_run(' are actively involved) coming from 78 universities in 28 countries.')

    p = document.add_paragraph('I am writing you this email since many people within ESTIEM are interested in continuing their career in a company such as ' + company_name +', we are eager to know what ')
    p.add_run('opportunities').bold = True
    p.add_run(' ' + company_name + ' could offer for Industrial Engineering and Management students. ')
    p.add_run(extra_reasons)

    document.add_paragraph('ESTIEM represents a professional network for Industrial Engineering and Management students, which means that our members share the same background in terms of studies, combining technological understanding with management skills. The aim of our organisation is to establish and foster international relations among European IEM students.  We have more than 30 international teams organising a high variety of Europe-wide activities such as exchanges, LSS courses, conferences, personal development, case study competitions, company visits and workshops. ')
    document.add_paragraph('The studies of IEM provide analytical capacities, engineering knowledge and practical management experiences, which make IEM students valuable since they are able to do business while understanding the underlying technology. ')
    document.add_paragraph('We see this opportunity as a way of cooperation with an international company that has interests in this field, such as yours, by presenting you with benefits that we, as a non-profit student organisation, can offer. ')
    document.add_paragraph('In the attachment of the mail, I am sending you our Company Brochure that contains all the detailed information about ESTIEM and how we could cooperate. If you are interested, I suggest that we schedule a meeting where we can present our goals and offers in more detail and consider a potential financial partnership. ')
    document.add_paragraph('If there is anything unclear or you\'d like more information, don\'t hesitate to contact us. \nWe appreciate your efforts and are looking forward to your answer. ')
    document.add_paragraph('Yours sincerely, \n')

    return document.save('Approach Email.docx')

def create_ln(greeting = 'Greetings,',
sender_name = 'EXAMPLE ESTIEMER',
sender_surname = 'zu ESTIEM',
company_name = 'INSERT COMPANY NAME', 
not_from_hr = False):
    
    document = Document()
    
    document.add_heading('LinkedIn Approach', 0)
    
    if not_from_hr:
      document.add_paragraph(greeting)
      
      document.add_paragraph('My name is ' + sender_name + ' and I am contacting you in the name of a student organisation called ESTIEM (estiem.org), the biggest European-wide student organisation for Industrial Engineering and Management students with approximately 70 000 students reached (among them, 10 000 are actively involved) coming from 78 universities in 28 countries. ')
      
      document.add_paragraph('I am writing you this message since many people within ESTIEM are interested in continuing their career in a company such as ' + company_name + ', we are eager to know what opportunities ' + company_name + ' could offer for Industrial Engineering and Management students.')
      
      document.add_paragraph('We see this opportunity as a way of cooperation with an international company that has interests in this field, such as yours, by presenting you with benefits that we, as a non-profit student organisation, can offer.')
      
      document.add_paragraph('We would appreciate it if you could redirect me to the right contact that is responsible for external relations in ' + company_name + '? We want to present you our brochure that contains more information about our organisation as well as the offerings.') 
      
      document.add_paragraph('Thank you for your time.')
      
      document.add_paragraph('Best regards,')
      
      document.add_paragraph(sender_name + ' ' + sender_surname)
    else:
        document.add_paragraph(greeting)
        
        document.add_paragraph('My name is ' + sender_name + ' and I am contacting you in the name of a student organisation called ESTIEM (estiem.org), the biggest European-wide student organisation for Industrial Engineering and Management students with approximately 70 000 students reached (among them, 10 000 are actively involved) coming from 78 universities in 28 countries.')
        
        document.add_paragraph('I am writing you this message since many people within ESTIEM are interested in continuing their career in a company such as ' + company_name + ', we are eager to know what opportunities ' + company_name + ' could offer for Industrial Engineering and Management students.')
        
        document.add_paragraph('We see this opportunity as a way of cooperation with an international company that has interests in this field, such as yours, by presenting you with benefits that we, as a non-profit student organisation, can offer.') 
        
        document.add_paragraph('We would like to get in contact with you and send you our brochure that contains more information about our organisation, as well as the offerings, if you are interested. If you are the right person whom I should contact, it would be my pleasure to send you this brochure to your email, and if this is not part of your job, can I ask you to give me any contact information of the person responsible with external relations?')
        
        document.add_paragraph('Thank you for taking the time to read this message.')
        
        document.add_paragraph('In the hope that we will achieve a successful cooperation,')
        
        document.add_paragraph('Best regards,')
        
        document.add_paragraph(sender_name + ' ' + sender_surname)
    
    return document.save('Approach LinkedIn Message.docx')
